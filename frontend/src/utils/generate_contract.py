from docx import Document
import re
from datetime import datetime

# Mock data from your system (in a real app, this would come from an API or database)
contract_data = {
    "id": "HD001",
    "start_date": "2025-01-01",
    "end_date": "2025-12-31",
    "status": "Đang hoạt động",
    "terms": "Thanh toán hàng tháng, không được chuyển nhượng.",
    "address": "123 Đường Láng, Đống Đa, Hà Nội",
    "duration": 12,
    "price": 3500000,
    "price_in_words": "Ba triệu năm trăm nghìn",
    "name_a": "Trần Thị B",
    "birth_year_a": "1980",
    "id_number_a": "123456789",
    "id_issue_date_a": "2010-05-15",
    "id_issue_place_a": "Hà Nội",
    "address_a": "456 Đường Giải Phóng, Hà Nội",
    "phone_a": "0987654321",
    "name_b": "Nguyễn Văn A",
    "birth_year_b": "1995",
    "id_number_b": "987654321",
    "id_issue_date_b": "2015-08-20",
    "id_issue_place_b": "Hà Nội",
    "address_b": "789 Đường Nguyễn Trãi, Hà Nội",
    "phone_b": "0901234567",
    "permanent_address_a": "456 Đường Giải Phóng, Hà Nội",
    "permanent_address_b": "789 Đường Nguyễn Trãi, Hà Nội",
    "payment_date": "5",
    "electricity_recipient": "Bên A",
    "electricity_payment_date": "5",
    "electricity_price": "3000 đ/kWh",
    "water_recipient": "Bên A",
    "water_payment_date": "5",
    "water_price": "20000 đ/m³",
    "other_fees": "Không có",
    "deposit": 3500000,
    "deposit_in_words": "Ba triệu năm trăm nghìn",
    "pages": 2
}

# Get current date and format it as "ngày DD tháng MM năm YYYY"
current_date = datetime.now()
formatted_date = f"ngày {current_date.day:02d} tháng {current_date.month:02d} năm {current_date.year}"
contract_data["current_date"] = formatted_date

def replace_placeholder(doc, placeholder, value):
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            for run in paragraph.runs:
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, str(value))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if placeholder in cell.text:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if placeholder in run.text:
                                run.text = run.text.replace(placeholder, str(value))

def generate_contract(template_path, output_path, data):
    doc = Document(template_path)
    
    # Replace placeholders with data
    for key, value in data.items():
        placeholder = "{{" + key + "}}"
        replace_placeholder(doc, placeholder, value)
    
    doc.save(output_path)
    print(f"Generated contract saved to {output_path}")

# Example usage
template_path = "contract_template.docx"
output_path = f"contract_{contract_data['id']}.docx"
generate_contract(template_path, output_path, contract_data)